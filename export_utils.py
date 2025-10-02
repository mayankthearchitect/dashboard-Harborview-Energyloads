"""
Export utilities for BIM Execution Plan
Generates PDF and DOCX documents from project data
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io


class BEPExporter:
    """Base class for BIM Execution Plan exports"""

    def __init__(self, project_data):
        self.data = project_data

    def format_list(self, items):
        """Format list items for display"""
        if isinstance(items, list):
            return ', '.join(items) if items else 'Not specified'
        return items if items else 'Not specified'

    def format_date(self, date_str):
        """Format date string"""
        if date_str:
            try:
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                return date_obj.strftime('%B %d, %Y')
            except:
                return date_str
        return 'Not specified'


class PDFExporter(BEPExporter):
    """PDF export functionality using ReportLab"""

    def generate(self, output_path=None):
        """Generate PDF document"""

        # Create buffer if no output path
        if output_path:
            buffer = output_path
        else:
            buffer = io.BytesIO()

        # Create document
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)

        # Container for document elements
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=10,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#475569'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        normal_style = styles['Normal']

        # Title Page
        elements.append(Spacer(1, 1.5*inch))
        elements.append(Paragraph("BIM EXECUTION PLAN", title_style))
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph(self.data.get('projectName', 'Untitled Project'), heading_style))
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph(f"Owner: {self.data.get('ownerName', 'N/A')}", normal_style))
        elements.append(Spacer(1, 0.1*inch))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", normal_style))
        elements.append(PageBreak())

        # Section 1: Project Information
        elements.append(Paragraph("1. PROJECT INFORMATION", heading_style))
        elements.append(Spacer(1, 0.1*inch))

        project_info = [
            ['Project Name:', self.data.get('projectName', 'N/A')],
            ['Owner:', self.data.get('ownerName', 'N/A')],
            ['Delivery Method:', self.data.get('deliveryMethod', 'N/A')],
            ['Facility Type:', self.data.get('facilityType', 'N/A')],
            ['Location:', self.data.get('projectLocation', 'N/A')],
            ['Project Value:', f"${self.data.get('projectValue', 0):,.0f}" if self.data.get('projectValue') else 'N/A'],
            ['Project Area:', f"{self.data.get('projectArea', 0):,.0f} sq ft" if self.data.get('projectArea') else 'N/A'],
            ['BIM Standard:', self.data.get('standardUsed', 'N/A')],
        ]

        t = Table(project_info, colWidths=[2*inch, 4.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.2*inch))

        elements.append(Paragraph("Project Description:", subheading_style))
        elements.append(Paragraph(self.data.get('projectDescription', 'N/A'), normal_style))

        # Section 2: Project Schedule
        elements.append(Paragraph("2. PROJECT SCHEDULE", heading_style))
        elements.append(Spacer(1, 0.1*inch))

        schedule_info = [
            ['Project Start:', self.format_date(self.data.get('projectStartDate'))],
            ['Project End:', self.format_date(self.data.get('projectEndDate'))],
            ['Design Phase End:', self.format_date(self.data.get('designPhaseEnd'))],
            ['Construction Start:', self.format_date(self.data.get('constructionStart'))],
        ]

        t = Table(schedule_info, colWidths=[2*inch, 4.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(t)

        if self.data.get('keyMilestones'):
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Key Milestones:", subheading_style))
            elements.append(Paragraph(self.data.get('keyMilestones'), normal_style))

        # Section 3: Contacts
        if self.data.get('contacts'):
            elements.append(Paragraph("3. KEY PROJECT CONTACTS", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            contact_data = [['Organization', 'Role', 'Name', 'Email']]
            for contact in self.data.get('contacts', []):
                contact_data.append([
                    contact.get('organization', 'N/A'),
                    contact.get('role', 'N/A'),
                    contact.get('contact_name', 'N/A'),
                    contact.get('email', 'N/A')
                ])

            t = Table(contact_data, colWidths=[1.5*inch, 1.3*inch, 1.5*inch, 2.2*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')])
            ]))
            elements.append(t)

        # Section 4: BIM Uses and Goals
        elements.append(Paragraph("4. BIM USES AND PROJECT GOALS", heading_style))
        elements.append(Spacer(1, 0.1*inch))

        if self.data.get('bim_uses'):
            elements.append(Paragraph("Selected BIM Uses:", subheading_style))
            bim_uses_text = self.format_list(self.data.get('bim_uses'))
            elements.append(Paragraph(bim_uses_text, normal_style))
            elements.append(Spacer(1, 0.1*inch))

        if self.data.get('projectGoals'):
            elements.append(Paragraph("Project Goals:", subheading_style))
            elements.append(Paragraph(self.data.get('projectGoals'), normal_style))

        # Section 5: BIM Roles
        if self.data.get('bim_roles'):
            elements.append(Paragraph("5. BIM ROLES AND RESPONSIBILITIES", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            roles = self.data.get('bim_roles', {})
            roles_info = [
                ['BIM Manager:', roles.get('bim_manager', 'Not specified')],
                ['BIM Coordinator:', roles.get('bim_coordinator', 'Not specified')],
            ]

            t = Table(roles_info, colWidths=[2*inch, 4.5*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(t)

            if roles.get('model_authors'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph("Model Authors:", subheading_style))
                elements.append(Paragraph(roles.get('model_authors'), normal_style))

        # Section 6: Collaboration
        if self.data.get('collaboration'):
            elements.append(Paragraph("6. COLLABORATION AND COMMUNICATION", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            collab = self.data.get('collaboration', {})
            collab_info = [
                ['Platform:', collab.get('collaboration_platform', 'N/A')],
                ['Meeting Schedule:', collab.get('meeting_schedule', 'N/A')],
                ['File Naming:', collab.get('file_naming_convention', 'N/A')],
            ]

            t = Table(collab_info, colWidths=[2*inch, 4.5*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(t)

        # Section 7: Technology
        if self.data.get('technology') or self.data.get('software'):
            elements.append(Paragraph("7. TECHNOLOGY REQUIREMENTS", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            if self.data.get('software'):
                elements.append(Paragraph("Software Applications:", subheading_style))
                software_text = self.format_list(self.data.get('software'))
                elements.append(Paragraph(software_text, normal_style))
                elements.append(Spacer(1, 0.1*inch))

            if self.data.get('technology'):
                tech = self.data.get('technology', {})
                tech_info = [
                    ['File Formats:', tech.get('file_formats', 'N/A')],
                    ['Coordinate System:', tech.get('coordinate_system', 'N/A')],
                ]

                t = Table(tech_info, colWidths=[2*inch, 4.5*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ]))
                elements.append(t)

        # Section 8: Model Management
        if self.data.get('model_management'):
            elements.append(Paragraph("8. MODEL MANAGEMENT", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            model = self.data.get('model_management', {})

            if model.get('lod_requirements'):
                elements.append(Paragraph("Level of Development (LOD) Requirements:", subheading_style))
                elements.append(Paragraph(model.get('lod_requirements'), normal_style))
                elements.append(Spacer(1, 0.1*inch))

            if model.get('model_structure'):
                elements.append(Paragraph("Model Structure:", subheading_style))
                elements.append(Paragraph(model.get('model_structure'), normal_style))

        # Section 9: Quality Control
        if self.data.get('quality_control'):
            elements.append(Paragraph("9. QUALITY CONTROL", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            qc = self.data.get('quality_control', {})

            qc_info = [
                ['Clash Detection:', qc.get('clash_detection_schedule', 'N/A')],
            ]

            t = Table(qc_info, colWidths=[2*inch, 4.5*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(t)

            if qc.get('qa_process'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph("QA Process:", subheading_style))
                elements.append(Paragraph(qc.get('qa_process'), normal_style))

        # Section 10: Deliverables
        if self.data.get('deliverables_risk'):
            elements.append(Paragraph("10. DELIVERABLES AND RISK MANAGEMENT", heading_style))
            elements.append(Spacer(1, 0.1*inch))

            deliv = self.data.get('deliverables_risk', {})

            if deliv.get('project_deliverables'):
                elements.append(Paragraph("Project Deliverables:", subheading_style))
                elements.append(Paragraph(deliv.get('project_deliverables'), normal_style))
                elements.append(Spacer(1, 0.1*inch))

            if deliv.get('risk_register'):
                elements.append(Paragraph("Risk Register:", subheading_style))
                elements.append(Paragraph(deliv.get('risk_register'), normal_style))

        # Build PDF
        doc.build(elements)

        if not output_path:
            buffer.seek(0)
            return buffer
        return output_path


class DOCXExporter(BEPExporter):
    """DOCX export functionality using python-docx"""

    def generate(self, output_path=None):
        """Generate DOCX document"""

        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title Page
        title = doc.add_heading('BIM EXECUTION PLAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(37, 99, 235)

        doc.add_paragraph()

        project_title = doc.add_heading(self.data.get('projectName', 'Untitled Project'), 1)
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        owner_para = doc.add_paragraph(f"Owner: {self.data.get('ownerName', 'N/A')}")
        owner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Section 1: Project Information
        doc.add_heading('1. PROJECT INFORMATION', 1)

        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'

        data = [
            ('Project Name:', self.data.get('projectName', 'N/A')),
            ('Owner:', self.data.get('ownerName', 'N/A')),
            ('Delivery Method:', self.data.get('deliveryMethod', 'N/A')),
            ('Facility Type:', self.data.get('facilityType', 'N/A')),
            ('Location:', self.data.get('projectLocation', 'N/A')),
            ('Project Value:', f"${self.data.get('projectValue', 0):,.0f}" if self.data.get('projectValue') else 'N/A'),
            ('Project Area:', f"{self.data.get('projectArea', 0):,.0f} sq ft" if self.data.get('projectArea') else 'N/A'),
            ('BIM Standard:', self.data.get('standardUsed', 'N/A')),
        ]

        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()
        doc.add_heading('Project Description', 2)
        doc.add_paragraph(self.data.get('projectDescription', 'N/A'))

        # Section 2: Project Schedule
        doc.add_heading('2. PROJECT SCHEDULE', 1)

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        schedule_data = [
            ('Project Start:', self.format_date(self.data.get('projectStartDate'))),
            ('Project End:', self.format_date(self.data.get('projectEndDate'))),
            ('Design Phase End:', self.format_date(self.data.get('designPhaseEnd'))),
            ('Construction Start:', self.format_date(self.data.get('constructionStart'))),
        ]

        for i, (label, value) in enumerate(schedule_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        if self.data.get('keyMilestones'):
            doc.add_paragraph()
            doc.add_heading('Key Milestones', 2)
            doc.add_paragraph(self.data.get('keyMilestones'))

        # Section 3: Contacts
        if self.data.get('contacts'):
            doc.add_heading('3. KEY PROJECT CONTACTS', 1)

            contacts = self.data.get('contacts', [])
            table = doc.add_table(rows=len(contacts) + 1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header
            headers = ['Organization', 'Role', 'Name', 'Email']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Data
            for i, contact in enumerate(contacts, 1):
                table.rows[i].cells[0].text = contact.get('organization', 'N/A')
                table.rows[i].cells[1].text = contact.get('role', 'N/A')
                table.rows[i].cells[2].text = contact.get('contact_name', 'N/A')
                table.rows[i].cells[3].text = contact.get('email', 'N/A')

        # Section 4: BIM Uses
        doc.add_heading('4. BIM USES AND PROJECT GOALS', 1)

        if self.data.get('bim_uses'):
            doc.add_heading('Selected BIM Uses', 2)
            doc.add_paragraph(self.format_list(self.data.get('bim_uses')))

        if self.data.get('projectGoals'):
            doc.add_heading('Project Goals', 2)
            doc.add_paragraph(self.data.get('projectGoals'))

        # Section 5: BIM Roles
        if self.data.get('bim_roles'):
            doc.add_heading('5. BIM ROLES AND RESPONSIBILITIES', 1)

            roles = self.data.get('bim_roles', {})
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = 'BIM Manager:'
            table.rows[0].cells[1].text = roles.get('bim_manager', 'Not specified')
            table.rows[1].cells[0].text = 'BIM Coordinator:'
            table.rows[1].cells[1].text = roles.get('bim_coordinator', 'Not specified')

            for row in table.rows:
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            if roles.get('model_authors'):
                doc.add_paragraph()
                doc.add_heading('Model Authors', 2)
                doc.add_paragraph(roles.get('model_authors'))

        # Section 6: Collaboration
        if self.data.get('collaboration'):
            doc.add_heading('6. COLLABORATION AND COMMUNICATION', 1)

            collab = self.data.get('collaboration', {})
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Light Grid Accent 1'

            collab_data = [
                ('Platform:', collab.get('collaboration_platform', 'N/A')),
                ('Meeting Schedule:', collab.get('meeting_schedule', 'N/A')),
                ('File Naming:', collab.get('file_naming_convention', 'N/A')),
            ]

            for i, (label, value) in enumerate(collab_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value
                table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        # Section 7: Technology
        if self.data.get('technology') or self.data.get('software'):
            doc.add_heading('7. TECHNOLOGY REQUIREMENTS', 1)

            if self.data.get('software'):
                doc.add_heading('Software Applications', 2)
                doc.add_paragraph(self.format_list(self.data.get('software')))

            if self.data.get('technology'):
                tech = self.data.get('technology', {})
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Light Grid Accent 1'

                table.rows[0].cells[0].text = 'File Formats:'
                table.rows[0].cells[1].text = tech.get('file_formats', 'N/A')
                table.rows[1].cells[0].text = 'Coordinate System:'
                table.rows[1].cells[1].text = tech.get('coordinate_system', 'N/A')

                for row in table.rows:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

        # Section 8: Model Management
        if self.data.get('model_management'):
            doc.add_heading('8. MODEL MANAGEMENT', 1)

            model = self.data.get('model_management', {})

            if model.get('lod_requirements'):
                doc.add_heading('Level of Development (LOD) Requirements', 2)
                doc.add_paragraph(model.get('lod_requirements'))

            if model.get('model_structure'):
                doc.add_heading('Model Structure', 2)
                doc.add_paragraph(model.get('model_structure'))

        # Section 9: Quality Control
        if self.data.get('quality_control'):
            doc.add_heading('9. QUALITY CONTROL', 1)

            qc = self.data.get('quality_control', {})

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = 'Clash Detection Schedule:'
            table.rows[0].cells[1].text = qc.get('clash_detection_schedule', 'N/A')
            table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

            if qc.get('qa_process'):
                doc.add_paragraph()
                doc.add_heading('QA Process', 2)
                doc.add_paragraph(qc.get('qa_process'))

        # Section 10: Deliverables
        if self.data.get('deliverables_risk'):
            doc.add_heading('10. DELIVERABLES AND RISK MANAGEMENT', 1)

            deliv = self.data.get('deliverables_risk', {})

            if deliv.get('project_deliverables'):
                doc.add_heading('Project Deliverables', 2)
                doc.add_paragraph(deliv.get('project_deliverables'))

            if deliv.get('risk_register'):
                doc.add_heading('Risk Register', 2)
                doc.add_paragraph(deliv.get('risk_register'))

        # Save document
        if output_path:
            doc.save(output_path)
            return output_path
        else:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
